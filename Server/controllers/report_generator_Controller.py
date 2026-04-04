"""
report_generator.py
────────────────────────────────────────────────────────────────────────────────
Two separate .docx report generators for an interview session:

  • generate_candidate_report(session_doc)  → bytes
      Shows the candidate their own performance:
        - Per-question: their answer, score, strengths, weaknesses,
          missing points, feedback, and improvement tips.
        - Personality traits with candidate-facing labels & feedback.
        - NO recruiter summary, NO HR decision, NO cheating detection.

  • generate_recruiter_report(session_doc)  → bytes
      Shows the recruiter a full picture of the candidate:
        - Per-question: transcript, score, strengths, weaknesses,
          missing points, evaluator feedback.
        - Technical skill assessment + recruiter summary.
        - Personality traits with HR-facing labels & insights + HR decision.
        - Full cheating detection: face auth incident log, eye gaze,
          phone detection.

FastAPI usage example:
    from report_generator import generate_candidate_report, generate_recruiter_report
    from fastapi.responses import StreamingResponse
    import io

    @router.get("/sessions/{session_id}/report/candidate")
    async def candidate_report(session_id: str):
        doc = await db.sessions.find_one({"_id": ObjectId(session_id)})
        return StreamingResponse(
            io.BytesIO(generate_candidate_report(doc)),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=candidate_report_{session_id}.docx"},
        )

    @router.get("/sessions/{session_id}/report/recruiter")
    async def recruiter_report(session_id: str):
        doc = await db.sessions.find_one({"_id": ObjectId(session_id)})
        return StreamingResponse(
            io.BytesIO(generate_recruiter_report(doc)),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=recruiter_report_{session_id}.docx"},
        )

Dependencies:
    pip install python-docx
"""

from __future__ import annotations
import json
import io
from datetime import datetime
from typing import Any

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Colour palette ───────────────────────────────────────────────────────────
C_PRIMARY     = RGBColor(0x1E, 0x3A, 0x5F)
C_ACCENT      = RGBColor(0x2E, 0x75, 0xB6)
C_WARN        = RGBColor(0xC0, 0x00, 0x00)
C_OK          = RGBColor(0x1F, 0x6B, 0x3A)
C_YELLOW      = RGBColor(0x7B, 0x5E, 0x00)
C_GRAY        = RGBColor(0x59, 0x59, 0x59)

C_WARN_BG     = "FFE9E9"
C_OK_BG       = "E6F4EC"
C_YELLOW_BG   = "FFF8DC"
C_LIGHT_BG    = "F2F2F2"
C_LIGHT_HEX   = "D6E4F0"
C_WHITE       = "FFFFFF"
C_PRIMARY_HEX = "1E3A5F"

DEFAULT_QUESTION_TITLES = [
    "Unity Render Pipelines",
    "Blender Modifier for Low-Poly Mesh",
    "Memory-Pool Allocators in Game Engines",
]


# ─── Low-level XML / style helpers ───────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def _set_col_width(cell, width_dxa: int):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_table_width(table, width_dxa: int):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(width_dxa))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _page_setup(doc: DocxDocument):
    for section in doc.sections:
        section.page_width   = Inches(8.5)
        section.page_height  = Inches(11)
        section.left_margin  = section.right_margin  = Inches(1)
        section.top_margin   = section.bottom_margin = Inches(1)


def _section_heading(doc: DocxDocument, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = C_PRIMARY
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), C_PRIMARY_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _sub_heading(doc: DocxDocument, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = C_ACCENT


def _label_value(doc: DocxDocument, label: str, value: str, value_color: RGBColor | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.color.rgb = C_PRIMARY
    r2 = p.add_run(value)
    r2.font.color.rgb = value_color or C_GRAY


def _color_box(doc: DocxDocument, lines: list[str], bg_hex: str, text_color: RGBColor):
    table = doc.add_table(rows=1, cols=1)
    _set_table_width(table, 9360)
    cell = table.cell(0, 0)
    _set_cell_bg(cell, bg_hex)
    _set_cell_margins(cell, 120, 120, 180, 180)
    _set_col_width(cell, 9360)
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.add_paragraph() if i > 0 else cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.size = Pt(10)
        r.font.color.rgb = text_color


def _two_col_table(doc: DocxDocument, rows: list[tuple], col1: int = 3200, col2: int = 6160):
    table = doc.add_table(rows=len(rows), cols=2)
    _set_table_width(table, col1 + col2)
    for i, row_data in enumerate(rows):
        label, value = row_data[0], row_data[1]
        bg = row_data[2] if len(row_data) > 2 else C_WHITE

        c1 = table.cell(i, 0)
        _set_cell_bg(c1, C_LIGHT_HEX)
        _set_cell_margins(c1)
        _set_col_width(c1, col1)
        c1.text = ""
        r = c1.paragraphs[0].add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = C_PRIMARY

        c2 = table.cell(i, 1)
        _set_cell_bg(c2, bg)
        _set_cell_margins(c2)
        _set_col_width(c2, col2)
        c2.text = ""
        r2 = c2.paragraphs[0].add_run(value)
        r2.font.size = Pt(10)
        r2.font.color.rgb = C_GRAY


def _bullet_list(doc: DocxDocument, items: list[str], color: RGBColor):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        if p.runs:
            p.runs[0].font.color.rgb = color


def _cover(doc: DocxDocument, title: str, subtitle: str, meta_rows: list[tuple]):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = C_PRIMARY

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_after = Pt(16)
    sr = s.add_run(subtitle)
    sr.italic = True
    sr.font.size = Pt(13)
    sr.font.color.rgb = C_ACCENT

    _two_col_table(doc, meta_rows, 2800, 6560)
    doc.add_page_break()


def _parse_date(raw) -> str:
    if isinstance(raw, dict):
        raw = raw.get("$date", "")
    try:
        return datetime.fromisoformat(str(raw).replace("Z", "+00:00")).strftime("%d %b %Y  %H:%M UTC")
    except Exception:
        return str(raw)


def _extract_meta(session_doc: dict) -> dict:
    return {
        "session_date": _parse_date(session_doc.get("session_date", "")),
        "candidate_id": session_doc.get("candidate_id", "N/A"),
        "session_id":   str(session_doc.get("_id", {}).get("$oid", session_doc.get("_id", "N/A"))),
        "is_mock":      session_doc.get("is_mock", False),
        "answers":      session_doc.get("answers", []),
        "tech":         session_doc.get("final_summary", {}).get("technical", {}),
        "integrity":    session_doc.get("final_summary", {}).get("integrity", {}),
        "personality":  session_doc.get("personality", {}),
        "phone_det":    session_doc.get("phone_detection", {}),
    }


# ─── Shared content blocks ────────────────────────────────────────────────────

def _block_technical(doc: DocxDocument, answers: list, tech: dict, qtitles: list[str], mode: str):
    """
    mode = "candidate" → shows improvement tips, no recruiter summary
    mode = "recruiter" → shows recruiter summary, no tips
    """
    _section_heading(doc, "Technical Evaluation")
    _label_value(doc, "Overall Score", f"{tech.get('final_score', 'N/A')} / 100")

    for idx, ans in enumerate(answers):
        ev     = ans.get("evaluation", {})
        score  = int(ev.get("score", 0))
        qtitle = qtitles[idx] if idx < len(qtitles) else f"Question {idx + 1}"
        sc     = C_WARN if score == 0 else (C_OK if score >= 70 else C_YELLOW)

        _sub_heading(doc, f"Q{idx + 1}: {qtitle}")

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        lbl = p.add_run("Your Answer:" if mode == "candidate" else "Candidate Answer:")
        lbl.bold = True
        lbl.font.color.rgb = C_PRIMARY
        transcript = ans.get("speech_to_text", {}).get("transcription", "(No transcription)")
        _color_box(doc, [transcript], C_LIGHT_BG, C_GRAY)

        _label_value(doc, "Score", f"{ev.get('score', 0)} / 100", sc)

        if ev.get("strengths"):
            p = doc.add_paragraph()
            p.add_run("Strengths:").bold = True
            p.runs[0].font.color.rgb = C_OK
            _bullet_list(doc, ev["strengths"], C_OK)

        if ev.get("weaknesses"):
            p = doc.add_paragraph()
            p.add_run("Weaknesses:").bold = True
            p.runs[0].font.color.rgb = C_WARN
            _bullet_list(doc, ev["weaknesses"], C_WARN)

        if ev.get("missing_points"):
            p = doc.add_paragraph()
            p.add_run("Missing Points:").bold = True
            p.runs[0].font.color.rgb = C_YELLOW
            _bullet_list(doc, ev["missing_points"], C_YELLOW)

        p = doc.add_paragraph()
        p.add_run("Feedback:").bold = True
        _color_box(doc, [ev.get("overall_feedback", "")], C_LIGHT_BG, C_GRAY)
        doc.add_paragraph()

    # Summary
    _section_heading(doc, "Technical Summary")

    if tech.get("overall_weaknesses"):
        p = doc.add_paragraph()
        p.add_run("Overall Weaknesses:").bold = True
        p.runs[0].font.color.rgb = C_WARN
        _bullet_list(doc, tech["overall_weaknesses"], C_WARN)

    skill_rows = [
        (k.replace("_", " "), v,
         C_WARN_BG if any(x in v.lower() for x in ["insufficient", "absent", "poor"]) else C_OK_BG)
        for k, v in tech.get("skill_assessment", {}).items()
    ]
    if skill_rows:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Skill Assessment:").bold = True
        _two_col_table(doc, skill_rows, 3200, 6160)

    if mode == "recruiter" and tech.get("summary_for_recruiter"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Recruiter Summary:").bold = True
        _color_box(doc, [tech["summary_for_recruiter"]], C_LIGHT_HEX, C_PRIMARY)

    if mode == "candidate" and tech.get("tips_for_candidate"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Tips for Improvement:").bold = True
        _color_box(doc, [tech["tips_for_candidate"]], C_OK_BG, C_OK)


def _block_personality(doc: DocxDocument, personality: dict, mode: str):
    """
    mode = "candidate" → candidate_feedback, candidate summary, recommendations
    mode = "recruiter" → hr_report, HR summary, HR decision badge
    """
    _section_heading(doc, "Personality Traits (Big Five)")
    traits = personality.get("overall", {}).get("traits", {})

    if traits:
        rows = []
        for trait, data in traits.items():
            pct     = round(data["score"] * 100)
            bar     = "█" * round(pct / 10) + "░" * (10 - round(pct / 10))
            insight = data.get("candidate_feedback") if mode == "candidate" else data.get("hr_report", "")
            rows.append((f"{trait}  {bar}  {pct}%", f"{data.get('label','')} — {insight}"))
        _two_col_table(doc, rows, 3200, 6160)

    doc.add_paragraph()
    dominant = ", ".join(personality.get("overall", {}).get("summary", {}).get("dominant_traits", []))
    _label_value(doc, "Dominant Traits", dominant)

    if mode == "candidate":
        summary_lines = personality.get("overall", {}).get("candidate_view", {}).get("summary", [])
        if summary_lines:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Your Personality Summary:").bold = True
            _color_box(doc, summary_lines, C_LIGHT_HEX, C_PRIMARY)
        recs = [d.get("recommendation", "") for d in traits.values() if d.get("recommendation")]
        if recs:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Recommendations for You:").bold = True
            _bullet_list(doc, recs, C_ACCENT)

    else:  # recruiter
        summary_lines = personality.get("overall", {}).get("hr_view", {}).get("summary", [])
        if summary_lines:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("HR Personality Summary:").bold = True
            _color_box(doc, summary_lines, C_LIGHT_HEX, C_PRIMARY)
        hr_decision = personality.get("overall", {}).get("hr_view", {}).get("decision", "N/A")
        dec_color   = C_OK    if hr_decision == "Consider" else C_WARN
        dec_bg      = C_OK_BG if hr_decision == "Consider" else C_WARN_BG
        doc.add_paragraph()
        _color_box(doc, [f"HR Decision: {hr_decision}"], dec_bg, dec_color)


def _block_cheating(doc: DocxDocument, face_auth: dict, eye_gaze: dict, phone_det: dict,
                    answers: list, qtitles: list[str]):
    """Full cheating/integrity section — recruiter only."""
    _section_heading(doc, "Integrity & Cheating Detection")

    is_suspected = face_auth.get("status") == "Suspected" or eye_gaze.get("status") == "High Alerts"
    _color_box(
        doc,
        ["⚠  CHEATING SUSPECTED" if is_suspected else "✓  No Integrity Issues Detected"],
        C_WARN_BG if is_suspected else C_OK_BG,
        C_WARN    if is_suspected else C_OK,
    )

    # Face auth
    doc.add_paragraph()
    _sub_heading(doc, "Face Authentication")
    fc = face_auth.get("counts", {})
    _two_col_table(doc, [
        ("Status",           face_auth.get("status", "N/A"), C_WARN_BG if is_suspected else C_OK_BG),
        ("Total Incidents",  str(face_auth.get("incidents_count", 0))),
        ("No Face Detected", str(fc.get("no_face", 0))),
        ("Multiple Faces",   str(fc.get("multiple_faces", 0))),
        ("Different Person", str(fc.get("different_person", 0))),
    ], 3200, 6160)

    if face_auth.get("incidents"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Incident Log:").bold = True
        inc_rows = []
        for inc in face_auth["incidents"]:
            ts = inc.get("timestamp", "")
            try:
                ts = datetime.fromisoformat(ts).strftime("%H:%M:%S")
            except Exception:
                pass
            alert = inc.get("alert_type", "").replace("_", " ").title()
            bg = (C_WARN_BG   if "different" in inc.get("alert_type", "") else
                  C_YELLOW_BG if "multi"     in inc.get("alert_type", "") else
                  C_LIGHT_BG)
            inc_rows.append((f"Frame {inc.get('frame_index','?')}  —  {ts}", alert, bg))
        _two_col_table(doc, inc_rows, 4200, 5160)

    # Eye gaze
    doc.add_paragraph()
    _sub_heading(doc, "Eye Gaze Monitoring")
    _two_col_table(doc, [
        ("Status",               eye_gaze.get("status", "N/A"), C_WARN_BG if is_suspected else C_OK_BG),
        ("Total Warnings",       str(eye_gaze.get("total_warnings", 0))),
        ("Total Alert Duration", f"{eye_gaze.get('total_duration', 0)} seconds"),
    ], 3200, 6160)

    # Phone detection
    doc.add_paragraph()
    _sub_heading(doc, "Phone Detection")
    for qid, data in phone_det.items():
        qidx = next(
            (i for i, a in enumerate(answers) if str(a.get("question_id", {}).get("$oid", "")) == qid), -1
        )
        qlabel = f"Q{qidx+1}: {qtitles[qidx]}" if 0 <= qidx < len(qtitles) else qid
        sb = C_WARN_BG if data.get("is_cheating") else C_OK_BG
        _two_col_table(doc, [
            ("Question",        qlabel),
            ("Status",          data.get("severity", "N/A"), sb),
            ("Phone Detected",  "Yes" if data.get("is_cheating") else "No"),
            ("Cheating Events", str(data.get("summary", {}).get("cheating_events_count", 0))),
            ("Video Duration",  f"{data.get('summary',{}).get('video_duration_sec', 0)}s"),
            ("Frames Sampled",  str(data.get("summary", {}).get("total_frames_sampled", 0))),
        ], 3200, 6160)
        doc.add_paragraph()


def _serialize(doc: DocxDocument) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── Public API ───────────────────────────────────────────────────────────────

def generate_candidate_report(
    session_doc: dict[str, Any],
    question_titles: list[str] | None = None,
) -> bytes:
    """
    Candidate-facing report.
    Includes: Q&A evaluation, skill assessment, improvement tips, personality (candidate view).
    Excludes: recruiter summary, HR decision, cheating detection.
    """
    doc     = DocxDocument()
    _page_setup(doc)
    m       = _extract_meta(session_doc)
    qtitles = question_titles or DEFAULT_QUESTION_TITLES
    label   = "Mock Interview" if m["is_mock"] else "Official Interview"

    _cover(doc, "Your Interview Performance Report", label, [
        ("Session Date", m["session_date"]),
        ("Session Type", label),
    ])

    _block_technical(doc, m["answers"], m["tech"], qtitles, mode="candidate")

    doc.add_page_break()

    _block_personality(doc, m["personality"], mode="candidate")

    return _serialize(doc)


def generate_recruiter_report(
    session_doc: dict[str, Any],
    question_titles: list[str] | None = None,
) -> bytes:
    """
    Recruiter-facing report.
    Includes: Q&A evaluation, skill assessment, recruiter summary,
              personality (HR view + decision), full cheating detection.
    Excludes: candidate improvement tips, candidate-facing personality text.
    """
    doc     = DocxDocument()
    _page_setup(doc)
    m       = _extract_meta(session_doc)
    qtitles = question_titles or DEFAULT_QUESTION_TITLES
    label   = "Mock Interview" if m["is_mock"] else "Official Interview"

    _cover(doc, "Candidate Interview Report", label, [
        ("Session Date", m["session_date"]),
        ("Candidate ID", m["candidate_id"]),
        ("Session ID",   m["session_id"]),
        ("Session Type", label),
    ])

    _block_technical(doc, m["answers"], m["tech"], qtitles, mode="recruiter")

    doc.add_page_break()

    _block_personality(doc, m["personality"], mode="recruiter")

    doc.add_page_break()

    _block_cheating(
        doc,
        face_auth = m["integrity"].get("face_auth", {}),
        eye_gaze  = m["integrity"].get("eye_gaze", {}),
        phone_det = m["phone_det"],
        answers   = m["answers"],
        qtitles   = qtitles,
    )

    return _serialize(doc)


# ─── CLI test ─────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    path = sys.argv[1] if len(sys.argv) > 1 else "session.txt"
    with open(path) as f:
        raw = json.load(f)

    cb = generate_candidate_report(raw)
    with open("candidate_report.docx", "wb") as f:
        f.write(cb)
    print(f"Candidate report → candidate_report.docx  ({len(cb):,} bytes)")

    rb = generate_recruiter_report(raw)
    with open("recruiter_report.docx", "wb") as f:
        f.write(rb)
    print(f"Recruiter report → recruiter_report.docx  ({len(rb):,} bytes)")